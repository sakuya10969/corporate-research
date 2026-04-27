"""PDF / Word ファイル生成"""

from __future__ import annotations

import io
from datetime import date

import markdown as md
from docx import Document
from docx.shared import Pt
from weasyprint import HTML

from src.db.models import AnalysisResult


def _build_html(result: AnalysisResult, company_name: str) -> str:
    """markdown_page を HTML に変換してスタイルを付与する"""
    body = md.markdown(result.markdown_page or "", extensions=["tables", "fenced_code"])
    return f"""<!DOCTYPE html>
<html lang="ja">
<head>
<meta charset="utf-8">
<style>
  @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+JP:wght@400;700&display=swap');
  body {{ font-family: 'Noto Sans JP', sans-serif; font-size: 11pt; line-height: 1.8; color: #1E293B; margin: 2cm; }}
  h1 {{ font-size: 18pt; color: #1E293B; border-bottom: 2px solid #2563EB; padding-bottom: 4px; }}
  h2 {{ font-size: 14pt; color: #2563EB; margin-top: 1.5em; }}
  h3 {{ font-size: 12pt; color: #1E293B; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
  th, td {{ border: 1px solid #CBD5E1; padding: 6px 10px; }}
  th {{ background: #F1F5F9; font-weight: bold; }}
  a {{ color: #2563EB; }}
  ul, ol {{ padding-left: 1.5em; }}
</style>
</head>
<body>
{body}
</body>
</html>"""


def generate_pdf(result: AnalysisResult) -> bytes:
    """WeasyPrint で PDF バイト列を生成する"""
    company_name = (result.structured or {}).get("company_profile", {}).get("name", "企業")
    html = _build_html(result, company_name)
    pdf_bytes = HTML(string=html).write_pdf()
    return pdf_bytes


def generate_docx(result: AnalysisResult) -> bytes:
    """python-docx で Word バイト列を生成する"""
    doc = Document()
    structured = result.structured or {}
    summary = result.summary or {}
    profile = structured.get("company_profile", {})
    company_name = profile.get("name", "企業分析レポート")
    analyzed_at = result.created_at.strftime("%Y-%m-%d") if result.created_at else date.today().isoformat()

    # タイトル
    doc.add_heading(f"{company_name} 企業分析レポート", 0)
    doc.add_paragraph(f"分析日: {analyzed_at}")

    # 企業プロフィール
    if profile:
        doc.add_heading("企業プロフィール", 1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "項目"
        hdr[1].text = "内容"
        for label, key in [("社名", "name"), ("設立", "founded"), ("代表者", "ceo"), ("所在地", "location"), ("従業員数", "employees"), ("資本金", "capital")]:
            val = profile.get(key, "")
            if val:
                row = table.add_row().cells
                row[0].text = label
                row[1].text = val

    # 企業概要
    if summary.get("overview"):
        doc.add_heading("企業概要", 1)
        doc.add_paragraph(summary["overview"])

    # 事業モデル
    if summary.get("business_model"):
        doc.add_heading("事業モデル", 1)
        doc.add_paragraph(summary["business_model"])

    # 事業領域
    domains = structured.get("business_domains", [])
    if domains:
        doc.add_heading("事業領域", 1)
        for d in domains:
            doc.add_paragraph(d, style="List Bullet")

    # プロダクト
    products = structured.get("products", [])
    if products:
        doc.add_heading("プロダクト・サービス", 1)
        for p in products:
            doc.add_paragraph(p, style="List Bullet")

    # 財務情報
    fin = structured.get("financials", {})
    if any(fin.get(k) for k in ("revenue", "operating_income", "net_income", "growth_rate")):
        doc.add_heading("財務情報", 1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "項目"
        hdr[1].text = "値"
        for label, key in [("売上高", "revenue"), ("営業利益", "operating_income"), ("純利益", "net_income"), ("成長率", "growth_rate")]:
            if fin.get(key):
                row = table.add_row().cells
                row[0].text = label
                row[1].text = fin[key]

    # SWOT
    swot = summary.get("swot", {})
    if swot:
        doc.add_heading("SWOT分析", 1)
        for label, key in [("強み", "strengths"), ("弱み", "weaknesses"), ("機会", "opportunities"), ("脅威", "threats")]:
            items = swot.get(key, [])
            if items:
                doc.add_heading(label, 2)
                for item in items:
                    doc.add_paragraph(item, style="List Bullet")

    # リスク
    risks = summary.get("risks", [])
    if risks:
        doc.add_heading("リスク要因", 1)
        for r in risks:
            doc.add_paragraph(r, style="List Bullet")

    # 競合
    competitors = summary.get("competitors", [])
    if competitors:
        doc.add_heading("競合企業（推定）", 1)
        doc.add_paragraph(", ".join(competitors))

    # 展望
    if summary.get("outlook"):
        doc.add_heading("今後の展望", 1)
        doc.add_paragraph(summary["outlook"])

    # ニュース
    news = structured.get("news", [])
    if news:
        doc.add_heading("ニュース", 1)
        for item in news:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item.get("title", ""))
            run.bold = True
            if item.get("date"):
                p.add_run(f"  ({item['date']})")
            if item.get("summary"):
                doc.add_paragraph(item["summary"])

    # 参照ソース
    sources = result.sources or []
    if sources:
        doc.add_heading("参照ソース", 1)
        for s in sources:
            title = s.get("title", "") if isinstance(s, dict) else getattr(s, "title", "")
            url = s.get("url", "") if isinstance(s, dict) else getattr(s, "url", "")
            doc.add_paragraph(f"{title}: {url}", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
